from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl.styles import Font
from django.core.mail import send_mail
from django.shortcuts import render,redirect,get_object_or_404
from django.contrib import messages,auth
from django.http import HttpResponse
from .models import *
from django.contrib.auth import authenticate, login, get_user_model
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.db.models import Sum
from datetime import datetime
from .models import Vehicule, Trajet, Maintenance, Depense

User = get_user_model()


def addTrajet(request):
    if request.method=="POST":
        vehicule = Vehicule.objects.get(id=request.POST['vehicule'])
        kilometrage_parcouru = request.POST['kilometrage_parcouru']
        consommation_carburant = request.POST['consommation_carburant']
        date_depart = request.POST['date_depart']
        date_arrivee = request.POST['date_arrivee']

        try:
            trajet = Trajet.objects.create(Utilisateur=request.user,vehicule=vehicule,kilometrage_parcouru=kilometrage_parcouru,consommation_carburant=consommation_carburant,date_depart=date_depart,date_arrivee=date_arrivee)
            trajet.save()
            messages.info(request, "Votre Trajet a Bien etre Enregistres")
        except Exception:
            messages.error(request, "Error Something Went Wrong")

        return redirect('/ajouter/Trajet')


def ajouterPageTrajet(request):
    vehicules = Vehicule.objects.all()
    return render(request,'addTrajet.html',{'vehicules':vehicules})

def getVehicule(request,pk):
    vehicule = get_object_or_404(Vehicule,id=pk)
    return render(request,'vehicule.hmtl', {'vehicule':vehicule})

def rapport_annuel_page(request):
    if request.method=="POST":
        date_str = request.POST['date']


        date = datetime.strptime(date_str, '%Y-%m')

        annee = date.year
        mois = date.month

        debut = datetime(int(annee), 1, 1)
        fin = datetime(int(annee) + 1, 1, 1)

        vehicules = Vehicule.objects.all()
        data = []

        for v in vehicules:
            trajets = Trajet.objects.filter(vehicule=v, date_depart__range=(debut, fin))
            maintenances = Maintenance.objects.filter(vehicule=v, date__range=(debut, fin))
            couts = Depense.objects.filter(vehicule=v, date_depense__range=(debut, fin))

            total_km = trajets.aggregate(Sum('kilometrage_parcouru'))['kilometrage_parcouru__sum'] or 0
            total_carburant = couts.filter(categorie='carburant').aggregate(Sum('montant'))['montant__sum'] or 0
            total_maintenance = maintenances.aggregate(Sum('cout'))['cout__sum'] or 0
            total_assurance = couts.filter(categorie='assurance').aggregate(Sum('montant'))['montant__sum'] or 0
            autres = couts.exclude(categorie__in=['carburant', 'assurance']).aggregate(Sum('montant'))['montant__sum'] or 0

            data.append({
                'vehicule': v,
                'km': total_km,
                'carburant': total_carburant,
                'maintenance': total_maintenance,
                'assurance': total_assurance,
                'autres': autres,
                'total': total_carburant + total_maintenance + total_assurance + autres,
            })

        return render(request, 'rapport_annuel.html', {
            'data': data,
            'annee': annee,
            'mois': mois
        })


def export_rapport_annuel_excel(request, annee):
    debut = datetime(int(annee), 1, 1)
    fin = datetime(int(annee) + 1, 1, 1)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Rapport {annee}"

    headers = ["Véhicule", "Km Parcourus", "Carburant (MRU)", "Maintenance (MRU)", "Assurance (MRU)", "Autres (MRU)", "Total (MRU)"]
    ws.append(headers)
    for col in ws[1]:
        col.font = Font(bold=True)

    for v in Vehicule.objects.all():
        trajets = Trajet.objects.filter(vehicule=v, date_depart__range=(debut, fin))
        maintenances = Maintenance.objects.filter(vehicule=v, date__range=(debut, fin))
        couts = Depense.objects.filter(vehicule=v, date_depense__range=(debut, fin))

        total_km = trajets.aggregate(Sum('kilometrage_parcouru'))['kilometrage_parcouru__sum'] or 0
        total_carburant = couts.filter(categorie='carburant').aggregate(Sum('montant'))['montant__sum'] or 0
        total_maintenance = maintenances.aggregate(Sum('cout'))['cout__sum'] or 0
        total_assurance = couts.filter(categorie='assurance').aggregate(Sum('montant'))['montant__sum'] or 0
        autres = couts.exclude(categorie__in=['carburant', 'assurance']).aggregate(Sum('montant'))['montant__sum'] or 0
        total = total_carburant + total_maintenance + total_assurance + autres

        ws.append([v.marque + " " + v.modele, total_km, total_carburant, total_maintenance, total_assurance, autres, total])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f"rapport_annuel_{annee}.xlsx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    wb.save(response)
    return response



def export_rapport_annuel_word(request, annee):
    debut = datetime(int(annee), 1, 1)
    fin = datetime(int(annee) + 1, 1, 1)

    document = Document()
    document.add_heading(f'Rapport Financier Annuel - {annee}', 0)

    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Véhicule'
    hdr_cells[1].text = 'Km'
    hdr_cells[2].text = 'Carburant (MRU)'
    hdr_cells[3].text = 'Maintenance (MRU)'
    hdr_cells[4].text = 'Assurance (MRU)'
    hdr_cells[5].text = 'Autres (MRU)'
    hdr_cells[6].text = 'Total (MRU)'

    for v in Vehicule.objects.all():
        trajets = Trajet.objects.filter(vehicule=v, date__range=(debut, fin))
        maintenances = Maintenance.objects.filter(vehicule=v, date__range=(debut, fin))
        couts = Depense.objects.filter(vehicule=v, date__range=(debut, fin))

        km = trajets.aggregate(Sum('kilometrage'))['kilometrage__sum'] or 0
        carburant = couts.filter(type='carburant').aggregate(Sum('montant'))['montant__sum'] or 0
        maintenance = maintenances.aggregate(Sum('cout'))['cout__sum'] or 0
        assurance = couts.filter(type='assurance').aggregate(Sum('montant'))['montant__sum'] or 0
        autres = couts.exclude(type__in=['carburant', 'assurance']).aggregate(Sum('montant'))['montant__sum'] or 0
        total = carburant + maintenance + assurance + autres

        row_cells = table.add_row().cells
        row_cells[0].text = str(v)
        row_cells[1].text = str(km)
        row_cells[2].text = f"{carburant:.2f}"
        row_cells[3].text = f"{maintenance:.2f}"
        row_cells[4].text = f"{assurance:.2f}"
        row_cells[5].text = f"{autres:.2f}"
        row_cells[6].text = f"{total:.2f}"

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"rapport_annuel_{annee}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    document.save(response)
    return response

@login_required
def home(request):
    notifications = Notification.objects.filter(whom = request.user)
    return render(request, 'home.html',{'notifications':notifications})

@login_required
def dashbord(request):
    notifications = Notification.objects.filter(whom = request.user)
    vehicules = Vehicule.objects.all()

    search = request.GET.get("search")

    # location = request.GET.get("location")
    brand = request.GET.get("brand")
    type = request.GET.getlist("type")
    color = request.GET.get("color")
    transmission = request.GET.getlist("transmission")
    fuel = request.GET.getlist("fuel")

    if search:
        vehicules = vehicules.filter(marque__icontains=search)
    # if location:
    #     print(location)
    #     vehicules = vehicules.filter(location=location)
    # if brand:
    #     print(brand)
    #     vehicules = vehicules.filter(brand=brand)
    if color:
        vehicules = vehicules.filter(color=color)
    if type:
        vehicules = vehicules.filter(type_vehicule__in=type)
    if transmission:
        vehicules = vehicules.filter(transmission__in=transmission)
    if fuel:
        vehicules = vehicules.filter(type_carburant__in=fuel)

    return render(request, 'dashbord.html', {'vehicules': vehicules,'notifications':notifications})

@login_required
def conducteurs(request):
    conducteurs = User.objects.filter(role="CONDUCTEUR")
    search = request.GET.get("search")
    if search:
        conducteurs = conducteurs.filter(
            Q(username__icontains=search) |
            Q(nom__icontains=search) |
            Q(prenom__icontains=search)
        )
    return render(request,'conducteurs.html',{"conducteurs":conducteurs})


@login_required
def getVehicule(request,pk):
    vehicule = get_object_or_404(Vehicule, id=pk)
    return render(request,'vehicule.html',{"vehicule":vehicule})

@login_required
def addVehicule(request):

    if request.method == "POST":

        vehicule = Vehicule.objects.create(
            image = request.FILES.get('image'),
            marque = request.POST['marque'],
            immatriculation= request.POST['immatriculation'],
            date_achat=  request.POST['date_achat'],
            kilometrage =  request.POST['kilometrage'],
            type_Vehiculeburant =  request.POST['type_Vehiculeburant'],
            modele = request.POST['modele']
        )

        vehicule.save()

        return redirect('/home')
    return render(request, '404.html')
@login_required
def modifieVehicule(request, pk):
    vehicule = get_object_or_404(Vehicule, id=pk)

    if request.method == "POST":
        vehicule.marque = request.POST['marque']
        vehicule.modele = request.POST['modele']
        vehicule.immatriculation = request.POST['immatrucilation']
        if request.POST['date_achat']:
            vehicule.date_achat = request.POST['date_achat']
        vehicule.kilometrage = request.POST['kilometrage']
        vehicule.type_vehicule = request.POST['type_carburant']
        vehicule.statut = request.POST['statut']

        if request.FILES.get('image'):
            vehicule.image = request.FILES['image']

        vehicule.save()
        return redirect('/dashbord')
    return render(request, '404.html')

@login_required
def deleteVehicule(request, pk):
    if request.method == "GET":
        vehicule = Vehicule.objects.get(id=pk)

        vehicule.delete()

        return render(request , 'vehicules.html')

@login_required
def addConducteur(request):
    
    if request.method == "POST":

        Utilisateur = Utilisateur.objects.create_user(
            nom = request.POST['nom'],
            prenom = request.POST['prenom'],
            numero_permis = request.POST['numero_permis'],
            date_validite_permis = request.POST['date_validite_permis'],
            image = request.FILES.get('image'),
        )

        Utilisateur.save()

        return redirect('/Utilisateurs')
    return render(request, '404.html')

@login_required
def modifieConducteur(request, pk):
    if request.method == "POST":
        user = get_object_or_404(User, id=pk)
        can_drive_value = request.POST.get("can_drive")
        user.canDrive = can_drive_value == "yes"
        user.save()
        return redirect(request.META.get("HTTP_REFERER", "/"))
        
    return render(request, '404.html')

@login_required
def deleteConducteur(request, pk):
    if request.method == "GET":
        Utilisateur = get_object_or_404(Utilisateur, id=pk)

        Utilisateur.delete()

        return redirect('/Utilisateurs')
    return render(request, '404.html')

def login(request):
    if request.user.is_authenticated:
        return redirect('/home') 
    return render(request , 'signIn.html')

def register(request):
    if request.user.is_authenticated:
        return redirect('/home') 
    return render(request , 'signUp.html')

def userLogin(request):
    if request.method == "POST":
        username = request.POST['username']
        password = request.POST['password']
        
        user = authenticate(request,username=username, password=password)
        

        if user is None:

            messages.error(request,"invalid credentials")
            return redirect('/auth/login')
        
        
        auth.login(request,user)

        return redirect('/home')
    
    return render(request, '404.html')

def userRegister(request):

    if request.method == "POST":
        username = request.POST['username']
        nom = request.POST['nom']
        prenom = request.POST['prenom']
        email = request.POST['email']
        password = request.POST['password']

        user = User.objects.filter(Q(username = username) or Q(email=email))

        if user.exists():
            messages.info(request, 'Username or email already used')
            
            return redirect('/auth/register')
        
        user = User.objects.create_user(username = username, email = email,nom =nom,prenom=prenom)
        user.set_password(password)

        user.save()

        auth.login(request, user)

        return redirect('/home')
    
    return redirect('/auth/register')

def userLogout(request):
    auth.logout(request)
    return redirect('/auth/login')

@login_required
def addPageVehicule(request):
    return render(request, 'addVehicule.html')

@login_required
def addPageUtilisateur(request):
    # Utilisateur 9assdi conducteur
    return render(request, 'addUtilisateur.html')


@login_required
def profile(request, pk):
    return render(request, 'profile.html', {'user' : request.user})


@login_required
def updateProfile(request):
    if request.method == "POST":
        user = User.objects.get(id=request.user.id)

        username = request.POST['username']
        nom = request.POST['nom']
        prenom = request.POST['prenom']
        phone = request.POST['phone']
        email = request.POST['email']
        image = request.FILES.get('image')
        


        user.username = username
        user.nom = nom
        user.prenom = prenom
        user.email = email
        user.phone = phone
        if image:
            user.image=image

        user.save()

    return redirect(f"/profile/{request.user.username}")


@login_required
def vehicule_list(request):
    # ha4i vue dashbord eli vih search w eli vih filter mli
    vehicules = Vehicule.objects.all()
    # layhi ngb8 2welen lwtat kamlin

    location = request.GET.get("location")
    brand = request.GET.get("brand")
    Vehicule_types = request.GET.getlist("type")
    color = request.GET.get("color")
    transmission = request.GET.getlist("transmission")
    fuel_type = request.GET.getlist("fuel")

    if location:
        vehicules = vehicules.filter(location=location)
    if brand:
        vehicules = vehicules.filter(brand=brand)
    if color:
        vehicules = vehicules.filter(color=color)
    if Vehicule_types:
        vehicules = vehicules.filter(type_vehicule__in=Vehicule_types)
    if transmission:
        vehicules = vehicules.filter(transmission__in=transmission)
    if fuel_type:
        vehicules = vehicules.filter(type_carburant__in=fuel_type)

    return render(request, "dashbord.html", {"vehicules": vehicules})


#Export Rapport annuel et mensuel en word et excel

def export_word(request, mois, annee):
    debut = datetime(int(annee), int(mois), 1)
    if int(mois) == 12:
        fin = datetime(int(annee) + 1, 1, 1)
    else:
        fin = datetime(int(annee), int(mois) + 1, 1)
    vehicules = Vehicule.objects.all()
    data = []

    for v in vehicules:
        trajets = Trajet.objects.filter(vehicule=v, date_depart__range=(debut, fin))
        maintenances = Maintenance.objects.filter(vehicule=v, date__range=(debut, fin))
        couts = Depense.objects.filter(vehicule=v, date_depense__range=(debut, fin))

        total_km = trajets.aggregate(Sum('kilometrage_parcouru'))['kilometrage_parcouru__sum'] or 0
        total_carburant = couts.filter(categorie='carburant').aggregate(Sum('montant'))['montant__sum'] or 0
        total_maintenance = maintenances.aggregate(Sum('cout'))['cout__sum'] or 0
        total_assurance = couts.filter(categorie='assurance').aggregate(Sum('montant'))['montant__sum'] or 0
        autres = couts.exclude(categorie__in=['carburant', 'assurance']).aggregate(Sum('montant'))['montant__sum'] or 0

        data.append({
            'vehicule': v.marque + " " + v.modele,
            'km': total_km,
            'carburant': total_carburant,
            'maintenance': total_maintenance,
            'assurance': total_assurance,
            'autres': autres,
            'total': total_carburant + total_maintenance + total_assurance + autres,
        })    
    total_general = sum(l['total'] for l in data)

    doc = Document()
    doc.add_heading(f"Rapport Financier - {mois}/{annee}", 0)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Véhicule'
    hdr_cells[1].text = 'Km'
    hdr_cells[2].text = 'Carburant (MRU)'
    hdr_cells[3].text = 'Maintenance (MRU)'
    hdr_cells[4].text = 'Assurance (MRU)'
    hdr_cells[5].text = 'Autres (MRU)'
    hdr_cells[6].text = 'Total (MRU)'

    for ligne in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(ligne['vehicule'])
        row_cells[1].text = str(ligne['km'])
        row_cells[2].text = str(ligne['carburant'])
        row_cells[3].text = str(ligne['maintenance'])
        row_cells[4].text = str(ligne['assurance'])
        row_cells[5].text = str(ligne['autres'])
        row_cells[6].text = str(ligne['total'])

    # Total général
    doc.add_paragraph(f"\nTotal général : {total_general} MRU")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"rapport_{mois}_{annee}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response


def export_excel(request, mois, annee):
    debut = datetime(int(annee), int(mois), 1)
    if int(mois) == 12:
        fin = datetime(int(annee) + 1, 1, 1)
    else:
        fin = datetime(int(annee), int(mois) + 1, 1)
    vehicules = Vehicule.objects.all()
    data = []

    for v in vehicules:
        trajets = Trajet.objects.filter(vehicule=v, date_depart__range=(debut, fin))
        maintenances = Maintenance.objects.filter(vehicule=v, date__range=(debut, fin))
        couts = Depense.objects.filter(vehicule=v, date_depense__range=(debut, fin))

        total_km = trajets.aggregate(Sum('kilometrage_parcouru'))['kilometrage_parcouru__sum'] or 0
        total_carburant = couts.filter(categorie='carburant').aggregate(Sum('montant'))['montant__sum'] or 0
        total_maintenance = maintenances.aggregate(Sum('cout'))['cout__sum'] or 0
        total_assurance = couts.filter(categorie='assurance').aggregate(Sum('montant'))['montant__sum'] or 0
        autres = couts.exclude(categorie__in=['carburant', 'assurance']).aggregate(Sum('montant'))['montant__sum'] or 0

        data.append({
            'vehicule': v.marque + " " + v.modele,
            'km': total_km,
            'carburant': total_carburant,
            'maintenance': total_maintenance,
            'assurance': total_assurance,
            'autres': autres,
            'total': total_carburant + total_maintenance + total_assurance + autres,
        })

    total_general = sum(l['total'] for l in data)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Rapport {mois}-{annee}"

    # En-têtes
    header = ['Véhicule', 'Km', 'Carburant (MRU)', 'Maintenance (MRU)', 'Assurance (MRU)', 'Autres (MRU)', 'Total (MRU)']
    ws.append(header)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    # Lignes de données
    for ligne in data:
        ws.append([
            ligne['vehicule'],
            ligne['km'],
            ligne['carburant'],
            ligne['maintenance'],
            ligne['assurance'],
            ligne['autres'],
            ligne['total'],
        ])

    # Total général
    ws.append([""] * 6 + [f"{total_general}"])

    # Réponse HTTP
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    filename = f"rapport_{mois}_{annee}.xlsx"
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    wb.save(response)
    return response


def rapport_mensuel(request, annee, mois):
    debut = datetime(int(annee), int(mois), 1)
    if int(mois) == 12:
        fin = datetime(int(annee) + 1, 1, 1)
    else:
        fin = datetime(int(annee), int(mois) + 1, 1)

    vehicules = Vehicule.objects.all()
    data = []

    for v in vehicules:
        trajets = Trajet.objects.filter(vehicule=v, date_depart__range=(debut, fin))
        maintenances = Maintenance.objects.filter(vehicule=v, date__range=(debut, fin))
        couts = Depense.objects.filter(vehicule=v, date_depense__range=(debut, fin))

        total_km = trajets.aggregate(Sum('kilometrage_parcouru'))['kilometrage_parcouru__sum'] or 0
        total_carburant = couts.filter(categorie='carburant').aggregate(Sum('montant'))['montant__sum'] or 0
        total_maintenance = maintenances.aggregate(Sum('cout'))['cout__sum'] or 0
        total_assurance = couts.filter(categorie='assurance').aggregate(Sum('montant'))['montant__sum'] or 0
        autres = couts.exclude(categorie__in=['carburant', 'assurance']).aggregate(Sum('montant'))['montant__sum'] or 0

        data.append({
            'vehicule': v.marque + " " + v.modele,
            'km': total_km,
            'carburant': total_carburant,
            'maintenance': total_maintenance,
            'assurance': total_assurance,
            'autres': autres,
            'total': total_carburant + total_maintenance + total_assurance + autres,
        })

    return render(request, 'rapport_mensuel.html', {'data': data, 'mois': mois, 'annee': annee})


def programmerMaintenance(request):
    vehicules= Vehicule.objects.all()
    return render(request,'programmerMaintenance.html', {'vehicules':vehicules})

def addMaintenance(request):
    if request.method == "POST":
        date = request.POST['date']
        depense = request.POST['depense']
        description = request.POST['description']
        type_maintenance = request.POST['type_maintenance']

        selectedVehicule = request.POST.get("vehicule")
        vehicule = Vehicule.objects.get(id=selectedVehicule)

        maintenance = Maintenance.objects.create(vehicule=vehicule,cout=depense,description=description,date=date,type_maintenance=type_maintenance)
        maintenance.save()
        return redirect('/dashbord')
    return render(request,'404.html')